from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file
from datetime import datetime, date
import json, os, io
from functools import wraps
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import requests
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY') or 'cartera-mora-secret-2024'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

# --- Config ---
# Usuario admin
ADMIN_USER = os.environ.get('ADMIN_USER', 'admin')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', os.environ.get('APP_PASSWORD', 'admin123'))
# Usuario viewer
VIEWER_USER = os.environ.get('VIEWER_USER', '')
VIEWER_PASSWORD = os.environ.get('VIEWER_PASSWORD', '')

EMAIL_DESTINO = os.environ.get('EMAIL_DESTINO', '')
RESEND_API_KEY = os.environ.get('RESEND_API_KEY', '')
EMAIL_FROM = os.environ.get('EMAIL_FROM', 'onboarding@resend.dev')

DATA_DIR = os.environ.get('RAILWAY_VOLUME_MOUNT_PATH', '.')
FILES_DIR = os.path.join(DATA_DIR, 'archivos')
DATA_FILE = os.path.join(DATA_DIR, 'cartera.json')
LOG_FILE = os.path.join(DATA_DIR, 'actividad.json')
EVENTOS_FILE = os.path.join(DATA_DIR, 'eventos.json')
os.makedirs(FILES_DIR, exist_ok=True)

ESTADOS = ['PREJUDICIAL', 'JUICIO', 'SENTENCIA', 'EJECUCION', 'COBRADO/CERRADO', 'MEDIACION', 'MEDIACION CON ACUERDO', 'ACUERDO EXTRAJUDICIAL']
PERSPECTIVAS = ['Alta', 'Media', 'Baja', 'Incobrable', '']

# --- Log de actividad ---
def log_actividad(accion, detalle='', cliente=''):
    try:
        logs = []
        if os.path.exists(LOG_FILE):
            with open(LOG_FILE, 'r', encoding='utf-8') as f:
                logs = json.load(f)
        logs.insert(0, {
            'fecha': datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
            'usuario': session.get('username', 'sistema'),
            'accion': accion,
            'cliente': cliente,
            'detalle': detalle
        })
        # Mantener solo los últimos 500 registros
        logs = logs[:500]
        with open(LOG_FILE, 'w', encoding='utf-8') as f:
            json.dump(logs, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

# --- Calculo de intereses ---
def calcular_interes_factura(monto, fecha_mora_str, tasa_anual):
    """Interés simple: monto * tasa_diaria * dias"""
    if not fecha_mora_str or not monto or not tasa_anual:
        return 0
    try:
        fecha_mora = datetime.strptime(fecha_mora_str, '%Y-%m-%d').date()
        hoy = date.today()
        dias = (hoy - fecha_mora).days
        if dias <= 0:
            return 0
        tasa_diaria = tasa_anual / 365 / 100
        return round(monto * tasa_diaria * dias)
    except Exception:
        return 0

def calcular_totales_cliente(cliente, tasa_anual):
    """Calcula monto_original e intereses sumando todas las facturas"""
    facturas = cliente.get('facturas', [])
    if not facturas:
        return cliente.get('monto_original', 0) or 0, cliente.get('intereses', 0) or 0
    monto_total = sum(f.get('monto', 0) or 0 for f in facturas)
    intereses_total = sum(calcular_interes_factura(f.get('monto', 0), f.get('fecha_mora', ''), tasa_anual) for f in facturas)
    return monto_total, intereses_total

# --- Persistencia ---
def load_data():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            # Migracion: agregar campos nuevos si no existen
            if 'tasa_bna' not in data:
                data['tasa_bna'] = 60.0
            if 'facturas' not in data:
                data['facturas'] = {}
            if 'pagos' not in data:
                data['pagos'] = {}
            return data
    return cargar_datos_iniciales()

def save_data(data):
    with open(DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def cargar_datos_iniciales():
    clientes = [
        {"id": 1, "razon_social": "SAVEDRA NATALIA", "monto_original": 9318447, "intereses": 0, "estado": "PREJUDICIAL", "sub_estado": "ACTUAL", "fecha_gestion": "", "observaciones": "", "perspectiva": "Alta"},
        {"id": 2, "razon_social": "TRANSPORTES LUNA", "monto_original": 1600000, "intereses": 0, "estado": "PREJUDICIAL", "sub_estado": "ACTUAL", "fecha_gestion": "", "observaciones": "", "perspectiva": "Alta"},
        {"id": 3, "razon_social": "LA MARTINA AGROSERVICIOS", "monto_original": 11837000, "intereses": 0, "estado": "JUICIO", "sub_estado": "SENTENCIA", "fecha_gestion": "", "observaciones": "", "perspectiva": "Media"},
        {"id": 4, "razon_social": "TRES DECIMA SAS", "monto_original": 17000000, "intereses": 0, "estado": "PREJUDICIAL", "sub_estado": "ACTUAL", "fecha_gestion": "", "observaciones": "", "perspectiva": "Alta"},
        {"id": 5, "razon_social": "LA COLO AGROPECUARIA", "monto_original": 46167077, "intereses": 0, "estado": "JUICIO", "sub_estado": "PENAL", "fecha_gestion": "", "observaciones": "", "perspectiva": "Incobrable"},
        {"id": 6, "razon_social": "KANCHEFF ALFREDO", "monto_original": 11747356, "intereses": 0, "estado": "JUICIO", "sub_estado": "PENAL", "fecha_gestion": "", "observaciones": "", "perspectiva": "Media"},
        {"id": 7, "razon_social": "INDUSTRIAS SANTA BARBARA", "monto_original": 0, "intereses": 0, "estado": "PREJUDICIAL", "sub_estado": "", "fecha_gestion": "", "observaciones": "23.000 LITROS", "perspectiva": ""},
        {"id": 8, "razon_social": "MARTINEZ SANTAMARINA SEBASTIAN", "monto_original": 31304298, "intereses": 0, "estado": "JUICIO", "sub_estado": "SENTENCIA", "fecha_gestion": "", "observaciones": "", "perspectiva": "Incobrable"},
        {"id": 9, "razon_social": "AGUABLANCA SAS", "monto_original": 31304298, "intereses": 0, "estado": "JUICIO", "sub_estado": "SENTENCIA", "fecha_gestion": "", "observaciones": "", "perspectiva": "Incobrable"},
        {"id": 10, "razon_social": "ARIEL GARCIA", "monto_original": 1500000, "intereses": 0, "estado": "MEDIACION", "sub_estado": "MEDIACION CON ACUERDO", "fecha_gestion": "", "observaciones": "", "perspectiva": "Baja"},
        {"id": 11, "razon_social": "NOBEN SRL", "monto_original": 71200000, "intereses": 0, "estado": "PREJUDICIAL", "sub_estado": "", "fecha_gestion": "", "observaciones": "", "perspectiva": ""},
        {"id": 12, "razon_social": "MONTEROS LEMON SAS", "monto_original": 69299410, "intereses": 0, "estado": "PREJUDICIAL", "sub_estado": "", "fecha_gestion": "", "observaciones": "", "perspectiva": ""},
        {"id": 13, "razon_social": "VISION EMPRESARIAL NOROESTE SRL", "monto_original": 27623401, "intereses": 0, "estado": "JUICIO", "sub_estado": "EJECUCION DE SENTENCIA", "fecha_gestion": "", "observaciones": "", "perspectiva": "Media"},
        {"id": 14, "razon_social": "NIEVAS NELSON", "monto_original": 3827682, "intereses": 0, "estado": "JUICIO", "sub_estado": "SENTENCIA", "fecha_gestion": "", "observaciones": "", "perspectiva": "Media"},
    ]
    data = {"clientes": clientes, "historial": {}, "facturas": {}, "pagos": {}, "tasa_bna": 60.0, "next_id": 15}
    save_data(data)
    return data

# --- Auth ---
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        if session.get('role') != 'admin':
            return jsonify({'error': 'Sin permisos'}), 403
        return f(*args, **kwargs)
    return decorated

def is_admin():
    return session.get('role') == 'admin'

# --- Email ---
def enviar_email_cambio_estado(cliente, estado_anterior, estado_nuevo):
    if not RESEND_API_KEY or not EMAIL_DESTINO:
        return
    try:
        fecha = datetime.now().strftime('%d/%m/%Y %H:%M')
        requests.post(
            'https://api.resend.com/emails',
            headers={'Authorization': f'Bearer {RESEND_API_KEY}', 'Content-Type': 'application/json'},
            json={
                'from': EMAIL_FROM,
                'to': [EMAIL_DESTINO],
                'subject': f'[Cartera Mora] Cambio de estado: {cliente["razon_social"]}',
                'html': f"""
                <h2>Cambio de estado en cartera</h2>
                <p><strong>Cliente:</strong> {cliente['razon_social']}</p>
                <p><strong>Estado anterior:</strong> {estado_anterior}</p>
                <p><strong>Estado nuevo:</strong> {estado_nuevo}</p>
                <p><strong>Fecha:</strong> {fecha}</p>
                """
            },
            timeout=10
        )
    except Exception:
        pass

# --- Rutas ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        usuario = request.form.get('usuario', '').strip()
        password = request.form.get('password', '').strip()
        if usuario == ADMIN_USER and password == ADMIN_PASSWORD:
            session['logged_in'] = True
            session['role'] = 'admin'
            session['username'] = usuario
            return redirect(url_for('home'))
        elif VIEWER_USER and usuario == VIEWER_USER and password == VIEWER_PASSWORD:
            session['logged_in'] = True
            session['role'] = 'viewer'
            session['username'] = usuario
            return redirect(url_for('home'))
        error = 'Usuario o contraseña incorrectos'
    return render_template('login.html', error=error)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/')
@login_required
def home():
    return render_template('home.html', role=session.get('role','admin'), username=session.get('username',''))

@app.route('/app')
@login_required
def index():
    data = load_data()
    return render_template('index.html', clientes=data['clientes'], estados=ESTADOS, perspectivas=PERSPECTIVAS, role=session.get('role','admin'), username=session.get('username',''))

@app.route('/api/config', methods=['GET'])
@login_required
def get_config():
    data = load_data()
    return jsonify({'tasa_bna': data.get('tasa_bna', 60.0)})

@app.route('/api/config', methods=['PUT'])
@login_required
def update_config():
    data = load_data()
    body = request.json
    if 'tasa_bna' in body:
        data['tasa_bna'] = float(body['tasa_bna'])
    save_data(data)
    return jsonify({'tasa_bna': data['tasa_bna']})

@app.route('/api/clientes', methods=['GET'])
@login_required
def get_clientes():
    data = load_data()
    tasa = data.get('tasa_bna', 60.0)
    clientes = []
    for c in data['clientes']:
        c2 = dict(c)
        facturas = data.get('facturas', {}).get(str(c['id']), [])
        c2['facturas'] = facturas
        if facturas:
            mo_facturas = sum(f.get('monto', 0) or 0 for f in facturas)
            int_facturas = sum(calcular_interes_factura(f.get('monto', 0), f.get('fecha_mora', ''), tasa) for f in facturas)
            c2['monto_original'] = mo_facturas
            c2['intereses'] = int_facturas
        clientes.append(c2)
    return jsonify(clientes)

@app.route('/api/clientes', methods=['POST'])
@login_required
def add_cliente():
    data = load_data()
    nuevo = request.json
    nuevo['id'] = data['next_id']
    nuevo.setdefault('monto_original', 0)
    nuevo.setdefault('intereses', 0)
    nuevo.setdefault('estado', 'PREJUDICIAL')
    nuevo.setdefault('sub_estado', '')
    nuevo.setdefault('fecha_gestion', '')
    nuevo.setdefault('observaciones', '')
    nuevo.setdefault('perspectiva', '')
    data['clientes'].append(nuevo)
    data['next_id'] += 1
    data['historial'][str(nuevo['id'])] = []
    data['facturas'][str(nuevo['id'])] = []
    save_data(data)
    log_actividad('Cliente creado', cliente=nuevo.get('razon_social',''))
    return jsonify(nuevo)

@app.route('/api/clientes/<int:cid>', methods=['PUT'])
@login_required
def update_cliente(cid):
    data = load_data()
    updates = request.json
    for i, c in enumerate(data['clientes']):
        if c['id'] == cid:
            estado_anterior = c.get('estado', '')
            estado_nuevo = updates.get('estado', estado_anterior)
            data['clientes'][i].update(updates)
            save_data(data)
            if estado_anterior != estado_nuevo:
                enviar_email_cambio_estado(data['clientes'][i], estado_anterior, estado_nuevo)
                log_actividad('Cambio de estado', f'{estado_anterior} → {estado_nuevo}', c.get('razon_social',''))
            else:
                campos = ', '.join(updates.keys())
                log_actividad('Edición de cliente', f'Campos: {campos}', c.get('razon_social',''))
            return jsonify(data['clientes'][i])
    return jsonify({'error': 'No encontrado'}), 404

@app.route('/api/clientes/<int:cid>', methods=['DELETE'])
@admin_required
def delete_cliente(cid):
    data = load_data()
    cliente = next((c for c in data['clientes'] if c['id'] == cid), None)
    data['clientes'] = [c for c in data['clientes'] if c['id'] != cid]
    data['historial'].pop(str(cid), None)
    data['facturas'].pop(str(cid), None)
    save_data(data)
    log_actividad('Cliente eliminado', cliente=cliente.get('razon_social','') if cliente else str(cid))
    return jsonify({'ok': True})

@app.route('/api/historial/<int:cid>', methods=['GET'])
@login_required
def get_historial(cid):
    data = load_data()
    return jsonify(data['historial'].get(str(cid), []))

@app.route('/api/historial/<int:cid>', methods=['POST'])
@login_required
def add_historial(cid):
    data = load_data()
    entrada = request.json
    entrada['fecha'] = datetime.now().strftime('%d/%m/%Y %H:%M')
    if str(cid) not in data['historial']:
        data['historial'][str(cid)] = []
    data['historial'][str(cid)].insert(0, entrada)
    for c in data['clientes']:
        if c['id'] == cid:
            c['fecha_gestion'] = datetime.now().strftime('%d/%m/%Y')
            break
    save_data(data)
    return jsonify(entrada)

# --- Facturas ---
@app.route('/api/facturas/<int:cid>', methods=['GET'])
@login_required
def get_facturas(cid):
    data = load_data()
    tasa = data.get('tasa_bna', 60.0)
    facturas = data.get('facturas', {}).get(str(cid), [])
    # Agregar interes calculado a cada factura
    for f in facturas:
        f['interes_calculado'] = calcular_interes_factura(f.get('monto', 0), f.get('fecha_mora', ''), tasa)
        f['total'] = (f.get('monto', 0) or 0) + f['interes_calculado']
    return jsonify(facturas)

@app.route('/api/facturas/<int:cid>', methods=['POST'])
@login_required
def add_factura(cid):
    data = load_data()
    tasa = data.get('tasa_bna', 60.0)
    factura = request.json
    factura['id'] = datetime.now().strftime('%Y%m%d%H%M%S%f')
    factura.setdefault('numero', '')
    factura.setdefault('monto', 0)
    factura.setdefault('fecha_mora', '')
    factura.setdefault('descripcion', '')
    if str(cid) not in data.get('facturas', {}):
        data.setdefault('facturas', {})[str(cid)] = []
    data['facturas'][str(cid)].append(factura)
    save_data(data)
    factura['interes_calculado'] = calcular_interes_factura(factura['monto'], factura['fecha_mora'], tasa)
    factura['total'] = factura['monto'] + factura['interes_calculado']
    return jsonify(factura)

@app.route('/api/facturas/<int:cid>/<fid>/fecha', methods=['PUT'])
@login_required
def update_factura_fecha(cid, fid):
    data = load_data()
    body = request.json
    facturas = data.get('facturas', {}).get(str(cid), [])
    for f in facturas:
        if f.get('id') == fid:
            f['fecha_mora'] = body.get('fecha_mora', '')
            break
    save_data(data)
    return jsonify({'ok': True})

@app.route('/api/facturas/<int:cid>/<fid>', methods=['DELETE'])
@admin_required
def delete_factura(cid, fid):
    data = load_data()
    tasa = data.get('tasa_bna', 60.0)
    facturas = data.get('facturas', {}).get(str(cid), [])
    data['facturas'][str(cid)] = [f for f in facturas if f.get('id') != fid]
    # Recalcular monto_original con las facturas restantes
    facturas_new = data['facturas'][str(cid)]
    mo = sum(f.get('monto', 0) or 0 for f in facturas_new)
    for c in data['clientes']:
        if c['id'] == cid:
            c['monto_original'] = mo
            break
    save_data(data)
    return jsonify({'ok': True})

# --- Pagos ---
@app.route('/api/pagos/<int:cid>', methods=['GET'])
@login_required
def get_pagos(cid):
    data = load_data()
    return jsonify(data.get('pagos', {}).get(str(cid), []))

@app.route('/api/pagos/<int:cid>', methods=['POST'])
@login_required
def add_pago(cid):
    data = load_data()
    pago = request.json
    pago['id'] = datetime.now().strftime('%Y%m%d%H%M%S%f')
    pago['fecha_registro'] = datetime.now().strftime('%d/%m/%Y')
    pago.setdefault('monto', 0)
    pago.setdefault('tipo', 'parcial')  # 'parcial', 'total', 'cheque'
    pago.setdefault('descuenta_de', 'total')
    pago.setdefault('observacion', '')
    pago.setdefault('fecha_cobro_cheque', '')  # Para cheques diferidos

    # Determinar si el cheque ya venció o es futuro
    if pago['tipo'] == 'cheque' and pago.get('fecha_cobro_cheque'):
        try:
            fecha_cheque = datetime.strptime(pago['fecha_cobro_cheque'], '%Y-%m-%d').date()
            pago['cheque_cobrado'] = fecha_cheque <= date.today()
        except:
            pago['cheque_cobrado'] = False
    
    # La fecha visible es la de registro
    pago['fecha'] = pago['fecha_registro']

    if 'pagos' not in data:
        data['pagos'] = {}
    if str(cid) not in data['pagos']:
        data['pagos'][str(cid)] = []
    data['pagos'][str(cid)].insert(0, pago)

    # Si es pago total → marcar como COBRADO/CERRADO
    if pago['tipo'] == 'total':
        for c in data['clientes']:
            if c['id'] == cid:
                c['estado'] = 'COBRADO/CERRADO'
                c['fecha_cobro'] = datetime.now().strftime('%d/%m/%Y')
                break

    save_data(data)
    return jsonify(pago)

@app.route('/api/pagos/<int:cid>/<pid>', methods=['DELETE'])
@admin_required
def delete_pago(cid, pid):
    data = load_data()
    pagos = data.get('pagos', {}).get(str(cid), [])
    data.setdefault('pagos', {})[str(cid)] = [p for p in pagos if p.get('id') != pid]
    save_data(data)
    return jsonify({'ok': True})

@app.route('/api/resumen', methods=['GET'])
@login_required
def get_resumen():
    data = load_data()
    tasa = data.get('tasa_bna', 60.0)
    hoy = date.today()
    resumen = {}
    todos = ESTADOS + ['MEDIACION', 'MEDIACION CON ACUERDO', 'ACUERDO EXTRAJUDICIAL']
    for estado in todos:
        cs = [c for c in data['clientes'] if c['estado'] == estado]
        monto_total = 0
        adeudado_total = 0
        pagado_total = 0
        cheques_pendientes = 0
        for c in cs:
            facturas = data.get('facturas', {}).get(str(c['id']), [])
            if facturas:
                mo = sum(f.get('monto', 0) or 0 for f in facturas)
                int_ = sum(calcular_interes_factura(f.get('monto', 0), f.get('fecha_mora', ''), tasa) for f in facturas)
            else:
                mo = c.get('monto_original', 0) or 0
                int_ = c.get('intereses', 0) or 0
            
            # Descontar pagos parciales y cheques cobrados
            pagos = data.get('pagos', {}).get(str(c['id']), [])
            pagado = 0
            cheque_pend = 0
            for p in pagos:
                if p.get('tipo') in ('parcial', 'total'):
                    pagado += p.get('monto', 0) or 0
                elif p.get('tipo') == 'cheque':
                    # Si la fecha del cheque ya pasó, lo cuenta como cobrado
                    try:
                        fc = datetime.strptime(p.get('fecha_cobro_cheque', ''), '%Y-%m-%d').date()
                        if fc <= hoy:
                            pagado += p.get('monto', 0) or 0
                        else:
                            cheque_pend += p.get('monto', 0) or 0
                    except:
                        pass

            monto_total += mo
            adeudado_neto = max(0, (mo + int_) - pagado)
            adeudado_total += adeudado_neto
            pagado_total += pagado
            cheques_pendientes += cheque_pend

        resumen[estado] = {
            'cantidad': len(cs),
            'monto_original': monto_total,
            'total_adeudado': adeudado_total,
            'pagado': pagado_total,
            'cheques_pendientes': cheques_pendientes
        }
    return jsonify(resumen)

# --- Eventos del calendario ---
def load_eventos():
    if os.path.exists(EVENTOS_FILE):
        with open(EVENTOS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_eventos(eventos):
    with open(EVENTOS_FILE, 'w', encoding='utf-8') as f:
        json.dump(eventos, f, ensure_ascii=False, indent=2)

@app.route('/api/eventos', methods=['GET'])
@login_required
def get_eventos():
    return jsonify(load_eventos())

@app.route('/api/eventos/<fecha>', methods=['POST'])
@login_required
def add_evento(fecha):
    eventos = load_eventos()
    nuevo = request.json
    nuevo['id'] = datetime.now().strftime('%Y%m%d%H%M%S%f')
    if fecha not in eventos:
        eventos[fecha] = []
    eventos[fecha].append(nuevo)
    save_eventos(eventos)
    log_actividad('Evento agregado', f"{nuevo.get('titulo','')} ({fecha})")
    return jsonify(nuevo)

@app.route('/api/eventos/<fecha>/<eid>', methods=['DELETE'])
@login_required
def delete_evento(fecha, eid):
    eventos = load_eventos()
    if fecha in eventos:
        eventos[fecha] = [e for e in eventos[fecha] if e.get('id') != eid]
        if not eventos[fecha]:
            del eventos[fecha]
    save_eventos(eventos)
    return jsonify({'ok': True})

# --- Chequeo diario de eventos y envio de emails ---
def enviar_email_evento(evento, fecha_str, cuando):
    if not RESEND_API_KEY or not EMAIL_DESTINO:
        return
    try:
        icono = {'audiencia':'⚖️','vencimiento':'⏰','reunion':'🤝','recordatorio':'🔔','pago':'💰','otro':'📌'}.get(evento.get('tipo'), '📌')
        cliente_txt = f"<p><strong>Cliente:</strong> {evento.get('cliente')}</p>" if evento.get('cliente') else ""
        requests.post(
            'https://api.resend.com/emails',
            headers={'Authorization': f'Bearer {RESEND_API_KEY}', 'Content-Type': 'application/json'},
            json={
                'from': EMAIL_FROM,
                'to': [EMAIL_DESTINO],
                'subject': f'[Cartera Mora] {icono} {evento.get("titulo")} — {cuando}',
                'html': f"""
                <h2>{icono} Recordatorio de evento</h2>
                <p><strong>{evento.get('titulo')}</strong></p>
                <p><strong>Fecha:</strong> {fecha_str}</p>
                <p><strong>Cuándo:</strong> {cuando}</p>
                {cliente_txt}
                """
            },
            timeout=10
        )
    except Exception:
        pass

@app.route('/api/eventos/chequear', methods=['GET', 'POST'])
def chequear_eventos():
    """Ruta para ser llamada diariamente (cron) que revisa eventos y cheques de hoy y mañana"""
    eventos = load_eventos()
    data = load_data()
    hoy = date.today()
    manana = date.fromordinal(hoy.toordinal() + 1)
    hoy_str = hoy.strftime('%Y-%m-%d')
    manana_str = manana.strftime('%Y-%m-%d')

    enviados = 0

    # Chequear eventos del calendario
    if hoy_str in eventos:
        for ev in eventos[hoy_str]:
            enviar_email_evento(ev, hoy.strftime('%d/%m/%Y'), 'Hoy')
            enviados += 1
    if manana_str in eventos:
        for ev in eventos[manana_str]:
            enviar_email_evento(ev, manana.strftime('%d/%m/%Y'), 'Mañana')
            enviados += 1

    # Chequear cheques diferidos que vencen hoy o mañana
    if RESEND_API_KEY and EMAIL_DESTINO:
        for c in data['clientes']:
            pagos = data.get('pagos', {}).get(str(c['id']), [])
            for p in pagos:
                if p.get('tipo') == 'cheque' and not p.get('cheque_cobrado'):
                    fecha_cheque = p.get('fecha_cobro_cheque', '')
                    if fecha_cheque in [hoy_str, manana_str]:
                        cuando = 'Hoy' if fecha_cheque == hoy_str else 'Mañana'
                        try:
                            requests.post(
                                'https://api.resend.com/emails',
                                headers={'Authorization': f'Bearer {RESEND_API_KEY}', 'Content-Type': 'application/json'},
                                json={
                                    'from': EMAIL_FROM,
                                    'to': [EMAIL_DESTINO],
                                    'subject': f'[Cartera Mora] 🧾 Cheque a cobrar — {c["razon_social"]} — {cuando}',
                                    'html': f"""
                                    <h2>🧾 Cheque a cobrar</h2>
                                    <p><strong>Cliente:</strong> {c['razon_social']}</p>
                                    <p><strong>Monto:</strong> $ {p['monto']:,.0f}</p>
                                    <p><strong>Fecha de cobro:</strong> {fecha_cheque.split('-')[::-1] and '/'.join(fecha_cheque.split('-')[::-1])}</p>
                                    <p><strong>Cuándo:</strong> {cuando}</p>
                                    {f"<p><strong>Observación:</strong> {p['observacion']}</p>" if p.get('observacion') else ''}
                                    """
                                },
                                timeout=10
                            )
                            enviados += 1
                            # Marcar cheque como cobrado si es hoy
                            if fecha_cheque == hoy_str:
                                p['cheque_cobrado'] = True
                        except Exception:
                            pass

    # Chequear fechas de acuerdo extrajudicial que vencen hoy o mañana
    if RESEND_API_KEY and EMAIL_DESTINO:
        for c in data['clientes']:
            fecha_acuerdo = c.get('fecha_acuerdo', '')
            if fecha_acuerdo and fecha_acuerdo in [hoy_str, manana_str]:
                cuando = 'Hoy' if fecha_acuerdo == hoy_str else 'Mañana'
                monto_acordado = c.get('monto_acuerdo', 0)
                condiciones = c.get('condiciones_acuerdo', '')
                try:
                    requests.post(
                        'https://api.resend.com/emails',
                        headers={'Authorization': f'Bearer {RESEND_API_KEY}', 'Content-Type': 'application/json'},
                        json={
                            'from': EMAIL_FROM,
                            'to': [EMAIL_DESTINO],
                            'subject': f'[Cartera Mora] 🤝 Vencimiento de acuerdo — {c["razon_social"]} — {cuando}',
                            'html': f"""
                            <h2>🤝 Fecha de pago de acuerdo extrajudicial</h2>
                            <p><strong>Cliente:</strong> {c['razon_social']}</p>
                            <p><strong>Fecha acordada:</strong> {'/'.join(fecha_acuerdo.split('-')[::-1])}</p>
                            <p><strong>Cuándo:</strong> {cuando}</p>
                            {f"<p><strong>Monto acordado:</strong> $ {monto_acordado:,.0f}</p>" if monto_acordado else ''}
                            {f"<p><strong>Condiciones:</strong> {condiciones}</p>" if condiciones else ''}
                            """
                        },
                        timeout=10
                    )
                    enviados += 1
                except Exception:
                    pass

    save_data(data)
    return jsonify({'ok': True, 'enviados': enviados, 'fecha_chequeo': hoy_str})

# --- Log de actividad ---
@app.route('/api/log', methods=['GET'])
@admin_required
def get_log():
    if not os.path.exists(LOG_FILE):
        return jsonify([])
    with open(LOG_FILE, 'r', encoding='utf-8') as f:
        return jsonify(json.load(f))

@app.route('/api/log', methods=['DELETE'])
@admin_required
def clear_log():
    if os.path.exists(LOG_FILE):
        os.remove(LOG_FILE)
    return jsonify({'ok': True})

# --- Backup ---
@app.route('/backup/descargar')
@admin_required
def backup_descargar():
    log_actividad('Backup descargado')
    return send_file(DATA_FILE, mimetype='application/json', as_attachment=True,
                     download_name=f'backup_cartera_{datetime.now().strftime("%Y%m%d_%H%M")}.json')

@app.route('/exportar/excel')
@login_required
def exportar_excel():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    data = load_data()
    tasa = data.get('tasa_bna', 60.0)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Clientes'

    headers = ['Cliente / Razón Social', 'CUIT / DNI', 'Monto Original', 'Estado',
               'Perspectiva de Cobro', 'Observaciones', 'Domicilio', 'Localidad', 'CP', 'Provincia']

    header_fill = PatternFill(start_color='1a2e4a', end_color='1a2e4a', fill_type='solid')
    header_font = Font(color='FFFFFF', bold=True, size=11)

    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')

    for i, c in enumerate(data['clientes'], start=2):
        # Calcular monto total de facturas
        facturas = data.get('facturas', {}).get(str(c['id']), [])
        if facturas:
            monto = sum(f.get('monto', 0) or 0 for f in facturas)
        else:
            monto = c.get('monto_original', 0) or 0

        ws.cell(row=i, column=1, value=c.get('razon_social', ''))
        ws.cell(row=i, column=2, value=c.get('cuit', ''))
        ws.cell(row=i, column=3, value=monto)
        ws.cell(row=i, column=4, value=c.get('estado', ''))
        ws.cell(row=i, column=5, value=c.get('perspectiva', ''))
        ws.cell(row=i, column=6, value=c.get('observaciones', ''))
        ws.cell(row=i, column=7, value=c.get('domicilio', ''))
        ws.cell(row=i, column=8, value=c.get('localidad', ''))
        ws.cell(row=i, column=9, value=c.get('cp', ''))
        ws.cell(row=i, column=10, value=c.get('provincia', ''))

        # Fila alternada
        if i % 2 == 0:
            for col in range(1, 11):
                ws.cell(row=i, column=col).fill = PatternFill(start_color='f0f4f8', end_color='f0f4f8', fill_type='solid')

    # Anchos de columna
    anchos = [32, 18, 18, 16, 18, 30, 25, 18, 8, 15]
    for i, ancho in enumerate(anchos, start=1):
        ws.column_dimensions[get_column_letter(i)].width = ancho

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fecha = datetime.now().strftime('%Y%m%d')
    return send_file(buf,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True,
                     download_name=f'cartera_mora_{fecha}.xlsx')

@app.route('/exportar/pdf/cliente/<int:cid>')
@login_required
def exportar_pdf_cliente(cid):
    data = load_data()
    tasa = data.get('tasa_bna', 60.0)
    cliente = next((c for c in data['clientes'] if c['id'] == cid), None)
    if not cliente:
        return 'Cliente no encontrado', 404

    facturas = data.get('facturas', {}).get(str(cid), [])
    historial = data.get('historial', {}).get(str(cid), [])

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    def fmt(v): return f"$ {v:,.0f}".replace(',', '.') if v else '$ 0'

    color_header = colors.HexColor('#1a2e4a')
    color_alt = colors.HexColor('#f0f4f8')

    titulo_style = ParagraphStyle('titulo', fontSize=16, textColor=color_header, fontName='Helvetica-Bold', spaceAfter=4)
    sub_style = ParagraphStyle('sub', fontSize=9, textColor=colors.HexColor('#64748b'), spaceAfter=16)
    h2_style = ParagraphStyle('h2', fontSize=11, textColor=color_header, fontName='Helvetica-Bold', spaceAfter=8, spaceBefore=14)
    normal = ParagraphStyle('normal', fontSize=9, leading=14)

    elements = []

    # Encabezado
    elements.append(Paragraph(f"REPORTE DE CLIENTE — CARTERA EN MORA", titulo_style))
    elements.append(Paragraph(f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')} hs · Tasa de interés: {tasa}% anual", sub_style))

    # Datos del cliente
    mo_total = sum(f.get('monto', 0) or 0 for f in facturas) if facturas else (cliente.get('monto_original', 0) or 0)
    int_total = sum(calcular_interes_factura(f.get('monto', 0), f.get('fecha_mora', ''), tasa) for f in facturas) if facturas else (cliente.get('intereses', 0) or 0)
    tot_total = mo_total + int_total

    datos = [
        ['Razón Social', cliente.get('razon_social', '')],
        ['Estado', cliente.get('estado', '')],
        ['Sub-estado', cliente.get('sub_estado', '') or '—'],
        ['Perspectiva de Cobro', cliente.get('perspectiva', '') or '—'],
        ['Última Gestión', cliente.get('fecha_gestion', '') or '—'],
        ['Observaciones', cliente.get('observaciones', '') or '—'],
        ['Capital Adeudado', fmt(mo_total)],
        ['Intereses acumulados', fmt(int_total)],
        ['TOTAL ADEUDADO', fmt(tot_total)],
    ]

    t_datos = Table(datos, colWidths=[5*cm, 11*cm])
    t_datos.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#64748b')),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#1a2e4a')),
        ('TEXTCOLOR', (0, -1), (-1, -1), colors.white),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, -1), (-1, -1), 10),
        ('ROWBACKGROUND', (0, 0), (-1, -2), [colors.white, color_alt]),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#e2e8f0')),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(t_datos)

    # Facturas
    if facturas:
        elements.append(Paragraph("Detalle de Facturas", h2_style))
        f_headers = ['N° Factura', 'Descripción', 'Fecha Mora', 'Monto', 'Interés al día', 'Total']
        f_filas = [f_headers]
        for f in facturas:
            monto = f.get('monto', 0) or 0
            int_ = calcular_interes_factura(monto, f.get('fecha_mora', ''), tasa)
            fecha_fmt = f.get('fecha_mora', '—')
            if fecha_fmt and fecha_fmt != '—':
                fecha_fmt = fecha_fmt.split('-')
                fecha_fmt = f"{fecha_fmt[2]}/{fecha_fmt[1]}/{fecha_fmt[0]}"
            f_filas.append([
                f.get('numero', '—') or '—',
                f.get('descripcion', '—') or '—',
                fecha_fmt,
                fmt(monto), fmt(int_), fmt(monto + int_)
            ])
        f_filas.append(['', 'TOTALES', '', fmt(mo_total), fmt(int_total), fmt(tot_total)])

        t_fact = Table(f_filas, colWidths=[2.5*cm, 4*cm, 2.5*cm, 2.8*cm, 2.8*cm, 2.8*cm])
        t_fact.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), color_header),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ROWBACKGROUND', (0, 1), (-1, -2), [colors.white, color_alt]),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#e2e8f0')),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#cbd5e1')),
            ('ALIGN', (3, 0), (-1, -1), 'RIGHT'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ]))
        elements.append(t_fact)

    # Historial
    if historial:
        elements.append(Paragraph("Historial de Gestiones", h2_style))
        h_filas = [['Fecha', 'Tipo', 'Descripción']]
        for h in historial:
            h_filas.append([h.get('fecha', ''), h.get('tipo', ''), Paragraph(h.get('nota', ''), ParagraphStyle('hn', fontSize=8, leading=10))])
        t_hist = Table(h_filas, colWidths=[3.5*cm, 4*cm, 9.9*cm])
        t_hist.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), color_header),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ROWBACKGROUND', (0, 1), (-1, -1), [colors.white, color_alt]),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#cbd5e1')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ]))
        elements.append(t_hist)

    doc.build(elements)
    buf.seek(0)
    nombre = cliente.get('razon_social', 'cliente').replace(' ', '_').lower()
    return send_file(buf, mimetype='application/pdf', as_attachment=True,
                     download_name=f'reporte_{nombre}_{datetime.now().strftime("%Y%m%d")}.pdf')

# --- Archivos adjuntos ---
@app.route('/api/archivos/<int:cid>', methods=['GET'])
@login_required
def get_archivos(cid):
    carpeta = os.path.join(FILES_DIR, str(cid))
    if not os.path.exists(carpeta):
        return jsonify([])
    archivos = []
    for fname in os.listdir(carpeta):
        fpath = os.path.join(carpeta, fname)
        archivos.append({
            'nombre': fname,
            'size': os.path.getsize(fpath),
            'fecha': datetime.fromtimestamp(os.path.getmtime(fpath)).strftime('%d/%m/%Y %H:%M')
        })
    archivos.sort(key=lambda x: x['fecha'], reverse=True)
    return jsonify(archivos)

@app.route('/api/archivos/<int:cid>', methods=['POST'])
@login_required
def upload_archivo(cid):
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': 'No filename'}), 400
    carpeta = os.path.join(FILES_DIR, str(cid))
    os.makedirs(carpeta, exist_ok=True)
    filename = secure_filename(f.filename)
    f.save(os.path.join(carpeta, filename))
    return jsonify({'ok': True, 'nombre': filename})

@app.route('/api/archivos/<int:cid>/<filename>', methods=['GET'])
@login_required
def download_archivo(cid, filename):
    carpeta = os.path.join(FILES_DIR, str(cid))
    return send_file(os.path.join(carpeta, secure_filename(filename)), as_attachment=True)

@app.route('/api/archivos/<int:cid>/<filename>', methods=['DELETE'])
@admin_required
def delete_archivo(cid, filename):
    carpeta = os.path.join(FILES_DIR, str(cid))
    fpath = os.path.join(carpeta, secure_filename(filename))
    if os.path.exists(fpath):
        os.remove(fpath)
    return jsonify({'ok': True})

# --- Importar desde Excel ---
@app.route('/api/importar', methods=['POST'])
@login_required
def importar_excel():
    try:
        import openpyxl
    except ImportError:
        return jsonify({'error': 'openpyxl no instalado'}), 500

    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    f = request.files['file']
    data = load_data()

    wb = openpyxl.load_workbook(f, data_only=True)
    ws = wb.active

    # Detectar fila de encabezados
    header_row = 1
    headers = []
    for i, row in enumerate(ws.iter_rows(min_row=1, max_row=10), start=1):
        vals = [str(cell.value or '').strip().lower().replace('\n', ' ') for cell in row]
        if any('razón social' in v or 'razon social' in v or 'cliente' in v for v in vals):
            header_row = i
            headers = vals
            break

    if not headers:
        return jsonify({'error': 'No se encontraron encabezados válidos'}), 400

    def get_col(row, names):
        for name in names:
            for i, h in enumerate(headers):
                if name in h:
                    val = row[i].value
                    return val
        return None

    # Normalizar perspectiva
    def normalizar_perspectiva(p):
        if not p: return ''
        p = str(p).strip().upper()
        mapa = {'ALTA': 'Alta', 'MEDIA': 'Media', 'BAJA': 'Baja', 'INCOBRABLE': 'Incobrable'}
        return mapa.get(p, p.capitalize())

    # Normalizar estado
    ESTADOS_VALIDOS = ['PREJUDICIAL', 'JUICIO', 'SENTENCIA', 'EJECUCION', 'MEDIACION', 'COBRADO/CERRADO']
    def normalizar_estado(e):
        if not e: return 'PREJUDICIAL'
        e = str(e).strip().upper()
        return e if e in ESTADOS_VALIDOS else 'PREJUDICIAL'

    agregados = 0
    actualizados = 0

    for row in ws.iter_rows(min_row=header_row + 1):
        razon = get_col(row, ['razón social', 'razon social', 'cliente / razón', 'cliente / razon', 'nombre', 'cliente'])
        if not razon:
            continue
        razon = str(razon).strip()
        if not razon:
            continue

        monto = float(get_col(row, ['monto', 'importe', 'deuda']) or 0)
        observaciones = str(get_col(row, ['observ', 'nota']) or '').strip()
        perspectiva = normalizar_perspectiva(get_col(row, ['perspectiva']))
        estado = normalizar_estado(get_col(row, ['estado']))
        cuit = str(get_col(row, ['cuit', 'dni']) or '').strip()
        domicilio = str(get_col(row, ['domicilio', 'direccion', 'dirección']) or '').strip()
        localidad = str(get_col(row, ['localidad']) or '').strip()
        cp = str(get_col(row, ['cp', 'codigo postal', 'código']) or '').strip()
        provincia = str(get_col(row, ['provincia']) or '').strip()

        # Verificar si ya existe
        idx_existente = next((i for i, c in enumerate(data['clientes']) if c['razon_social'].strip().upper() == razon.upper()), None)

        if idx_existente is not None:
            c = data['clientes'][idx_existente]
            cid = c['id']
            actualizado = False

            # Actualizar campos vacíos
            for campo, valor in [('cuit', cuit), ('domicilio', domicilio), ('localidad', localidad),
                                  ('cp', cp), ('provincia', provincia), ('perspectiva', perspectiva)]:
                if valor and not c.get(campo):
                    data['clientes'][idx_existente][campo] = valor
                    actualizado = True

            # Observaciones: agregar si hay nuevas
            if observaciones and observaciones not in (c.get('observaciones') or ''):
                obs_actual = c.get('observaciones') or ''
                data['clientes'][idx_existente]['observaciones'] = (obs_actual + ' | ' + observaciones).strip(' |') if obs_actual else observaciones
                actualizado = True

            # Si hay monto nuevo, agregar como factura
            if monto > 0:
                if str(cid) not in data.get('facturas', {}):
                    data.setdefault('facturas', {})[str(cid)] = []
                nueva_factura = {
                    'id': datetime.now().strftime('%Y%m%d%H%M%S%f') + str(cid),
                    'numero': '',
                    'monto': monto,
                    'fecha_mora': '',
                    'descripcion': f'Importado desde Excel {datetime.now().strftime("%d/%m/%Y")}'
                }
                data['facturas'][str(cid)].append(nueva_factura)
                actualizado = True

            if actualizado:
                actualizados += 1
            continue

        # Cliente nuevo
        nuevo = {
            'id': data['next_id'],
            'razon_social': razon,
            'cuit': cuit,
            'domicilio': domicilio,
            'localidad': localidad,
            'cp': cp,
            'provincia': provincia,
            'monto_original': monto,
            'intereses': 0,
            'estado': estado,
            'sub_estado': '',
            'fecha_gestion': '',
            'observaciones': observaciones,
            'perspectiva': perspectiva,
        }
        data['clientes'].append(nuevo)
        data['historial'][str(nuevo['id'])] = []
        data.setdefault('facturas', {})[str(nuevo['id'])] = []

        # Si tiene monto, crear factura
        if monto > 0:
            data['facturas'][str(nuevo['id'])].append({
                'id': datetime.now().strftime('%Y%m%d%H%M%S%f') + str(nuevo['id']),
                'numero': '',
                'monto': monto,
                'fecha_mora': '',
                'descripcion': f'Importado desde Excel {datetime.now().strftime("%d/%m/%Y")}'
            })

        data['next_id'] += 1
        agregados += 1

    save_data(data)
    return jsonify({'ok': True, 'agregados': agregados, 'actualizados': actualizados, 'total': len(data['clientes'])})

@app.route('/exportar/carta/<int:cid>')
@login_required
def exportar_carta(cid):
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = load_data()
    tasa = data.get('tasa_bna', 60.0)
    cliente = next((c for c in data['clientes'] if c['id'] == cid), None)
    if not cliente:
        return 'Cliente no encontrado', 404

    facturas = data.get('facturas', {}).get(str(cid), [])
    hoy = datetime.now()
    meses = ['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre']
    lugar_fecha = f"Los Nogales, {meses[hoy.month-1]} de {hoy.year}."

    mo_total = sum(f.get('monto', 0) or 0 for f in facturas)
    int_total = sum(calcular_interes_factura(f.get('monto', 0), f.get('fecha_mora', ''), tasa) for f in facturas)
    tot_total = mo_total + int_total

    def fmt(v): return f"$ {v:,.0f}".replace(',', '.')

    nums_facturas = [f.get('numero', '') for f in facturas if f.get('numero', '')]
    if nums_facturas:
        if len(nums_facturas) == 1:
            ref_facturas = f"la factura Nº {nums_facturas[0]}"
        else:
            ref_facturas = f"las facturas Nº {', '.join(nums_facturas[:-1])} y Nº {nums_facturas[-1]}"
    else:
        ref_facturas = "las facturas impagas"

    razon_social = cliente.get('razon_social', '')
    cuit_cliente = cliente.get('cuit', '')

    doc = DocxDocument()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    def set_font(run, bold=False, size=11):
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.bold = bold

    # Cabecera — tabla dos columnas
    domicilio_cliente = cliente.get('domicilio', '')
    localidad_cliente = cliente.get('localidad', '')
    cp_cliente = cliente.get('cp', '')
    provincia_cliente = cliente.get('provincia', '')

    tabla_cab = doc.add_table(rows=5, cols=2)
    tabla_cab.style = 'Table Grid'
    cab_izq = ['LUBRE S.R.L.', 'RUTA NACIONAL Nº 9 KM 1306', 'LOS NOGALES', '4101', 'TUCUMÁN']
    cab_der = [razon_social, domicilio_cliente or '', localidad_cliente or '', cp_cliente or '', provincia_cliente or '']

    for i, (izq, der) in enumerate(zip(cab_izq, cab_der)):
        fila = tabla_cab.rows[i]
        c_izq = fila.cells[0]
        c_der = fila.cells[1]
        p_izq = c_izq.paragraphs[0]
        p_der = c_der.paragraphs[0]
        r_izq = p_izq.add_run(izq)
        r_der = p_der.add_run(der)
        bold_row = i < 2
        set_font(r_izq, bold=bold_row, size=10)
        set_font(r_der, bold=bold_row, size=10)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        for cell in [c_izq, c_der]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top','left','bottom','right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    doc.add_paragraph()

    # Lugar y fecha
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_fecha.add_run(lugar_fecha)
    set_font(r, bold=True, size=11)

    doc.add_paragraph()

    # Cuerpo
    cuit_txt = f", CUIT {cuit_cliente}," if cuit_cliente else ","
    dom_txt = f" con domicilio en {domicilio_cliente}," if domicilio_cliente else ","
    monto_txt = f" por un monto total de {fmt(tot_total)} (capital {fmt(mo_total)} más intereses {fmt(int_total)})" if tot_total else ""

    p_cuerpo = doc.add_paragraph()
    p_cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_run(p, text, bold=False):
        r = p.add_run(text)
        set_font(r, bold=bold, size=11)
        return r

    add_run(p_cuerpo, 'MARTÍN VEGA', bold=True)
    add_run(p_cuerpo, ', abogado de la matrícula, en mi carácter de apoderado legal de la firma ')
    add_run(p_cuerpo, 'LUBRE S.R.L.', bold=True)
    add_run(p_cuerpo, ' CUIT Nº 30-71005185-9, conforme mandato otorgado mediante escritura pública Nº 770 de fecha 26/09/2019, pasada ante la Escribanía Nicolás Federico Odstrcil, adscripto al Registro Notarial Nº 51, me dirijo a Ud. ')
    add_run(p_cuerpo, razon_social, bold=True)
    add_run(p_cuerpo, f'{cuit_txt}{dom_txt} en razón de no haber cancelado la totalidad de la deuda que tiene con mi mandante, emergente de {ref_facturas}{monto_txt}. Lo íntimo, en un plazo de 72 hs, al pago de las mismas, con más sus intereses, bajo apercibimiento de iniciar acción judicial que corresponde en vuestra contra. Pongo en su conocimiento que la cancelación deberá hacerla en el domicilio de LUBRE S.R.L., Ruta Nacional Nº 9, KM 1306, en el horario de 08:00 a 17:00 hs, Los Nogales, Teléfono Celular (381) 156069919. ')
    add_run(p_cuerpo, 'Queda Ud. debidamente intimado y notificado.', bold=True)
    add_run(p_cuerpo, '—' * 60)

    doc.add_paragraph()
    doc.add_paragraph()

    # Firma
    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_firma, '_______________________________')

    p_nombre = doc.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_nombre, 'MARTÍN VEGA', bold=True)

    p_cargo = doc.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_cargo, 'Abogado – Apoderado LUBRE S.R.L.')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    nombre = razon_social.replace(' ', '_').lower()
    return send_file(buf,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True,
                     download_name=f'carta_documento_{nombre}_{hoy.strftime("%Y%m%d")}.docx')

@app.route('/exportar/pdf')
@login_required
def exportar_pdf():
    data = load_data()
    tasa = data.get('tasa_bna', 60.0)
    clientes = data['clientes']

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    titulo_style = ParagraphStyle('titulo', parent=styles['Title'],
                                   fontSize=16, textColor=colors.HexColor('#1a2e4a'), spaceAfter=4)
    subtitulo_style = ParagraphStyle('sub', parent=styles['Normal'],
                                      fontSize=9, textColor=colors.HexColor('#64748b'), spaceAfter=12)
    cell_style = ParagraphStyle('cell', fontSize=7.5, leading=10)

    def fmt(v): return f"$ {v:,.0f}".replace(',', '.') if v else '-'

    elements = []
    elements.append(Paragraph("INFORME DE ESTADO DE DEUDA — CARTERA EN MORA", titulo_style))
    elements.append(Paragraph(f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')} hs · Tasa BNA: {tasa}% anual · Montos en ARS", subtitulo_style))

    headers = ['N°', 'Cliente / Razón Social', 'Capital Adeudado', 'Intereses', 'Total Adeudado', 'Estado', 'Última Gestión', 'Perspectiva', 'Observaciones']
    filas = [headers]
    total_mo = total_int = 0
    for i, c in enumerate(clientes):
        facturas = data.get('facturas', {}).get(str(c['id']), [])
        if facturas:
            mo, int_ = calcular_totales_cliente({'facturas': facturas}, tasa)
        else:
            mo = c.get('monto_original', 0) or 0
            int_ = c.get('intereses', 0) or 0
        total = mo + int_
        total_mo += mo
        total_int += int_
        filas.append([
            str(i + 1),
            Paragraph(c.get('razon_social', ''), cell_style),
            fmt(mo), fmt(int_), fmt(total),
            c.get('estado', ''),
            c.get('fecha_gestion', '') or '-',
            c.get('perspectiva', '') or '-',
            Paragraph(c.get('observaciones', '') or '-', cell_style),
        ])

    filas.append(['', 'TOTALES', fmt(total_mo), fmt(total_int), fmt(total_mo + total_int), '', '', '', ''])

    col_widths = [1*cm, 5.5*cm, 3*cm, 2.5*cm, 3*cm, 3*cm, 2.8*cm, 2.2*cm, 4.5*cm]
    tabla = Table(filas, colWidths=col_widths, repeatRows=1)
    color_header = colors.HexColor('#1a2e4a')
    color_alt = colors.HexColor('#f0f4f8')
    tabla.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), color_header),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 7.5),
        ('ROWBACKGROUND', (0, 1), (-1, -2), [colors.white, color_alt]),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#e2e8f0')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#cbd5e1')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('ALIGN', (2, 1), (4, -1), 'RIGHT'),
    ]))
    elements.append(tabla)

    doc.build(elements)
    buf.seek(0)
    return send_file(buf, mimetype='application/pdf', as_attachment=True,
                     download_name=f'cartera_mora_{datetime.now().strftime("%Y%m%d")}.pdf')

if __name__ == '__main__':
    app.run(debug=True)
